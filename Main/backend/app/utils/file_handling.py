"""File-format text extraction for resumes (PDF, DOCX).

Ported from Code/app_logic.py — behavior preserved verbatim. Phase 1 text
cleaning improvements are layered on top via text_cleaning.normalize_document.
"""

from __future__ import annotations

import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF file using PyMuPDF."""
    text = ""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
    except Exception as e:
        logger.error("PDF extraction error: %s", e)
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    text = ""
    try:
        import docx
        document = docx.Document(file_path)
        para_lines = [para.text for para in document.paragraphs if para.text.strip()]
        text += "\n".join(para_lines) + "\n"
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text += cell_text + "\n"
    except Exception as e:
        logger.error("DOCX extraction error: %s", e)
    return text
